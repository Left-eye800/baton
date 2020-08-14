"""将题目写入docx文件中

提供接口write_in_docx将dict写入docx文件之中

the statistics of this file:
lines(count)    understand_level(h/m/l)    classes(count)    functions(count)    fields(count)
000000000149    ----------------------h    00000000000000    0000000000000001    ~~~~~~~~~~~~2
"""
import pprint
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_json import read_from_json_file

__author__ = '与C同行'
try:
    json_filename = sys.argv[1]
    docx_path = sys.argv[2]
except IndexError:
    json_filename = r'.\source_files\json_source\example.json'
    docx_path = r'.\example.docx'


def write_in_docx(text_dict_para, docx_filename, paragraph_apace_after_para=Pt(10),
                  doc_font_name_para='Times New Roman', doc_font_size_para=Pt(12),
                  option_length_para=70):
    """将dict写入docx,写入过程可自行探索,不再赘述

    :param option_length_para: 选择题选项默认对齐长度
    :param doc_font_size_para: docx文档字体大小
    :param paragraph_apace_after_para: docx文档段落间距
    :param doc_font_name_para: docx文档字体
    :param text_dict_para: dict字典内容
    :param docx_filename: docx文件名
    """
    def write_question(paragraph_para, content_para):
        question = content_para['question']
        if question['type'] == 'text':
            question_text_list = question['content'].split('__')
            for slice_text in question_text_list:
                if slice_text.startswith('^'):
                    slice_text = slice_text[1:]
                    paragraph_para.add_run(' ').underline = True
                    paragraph_para.add_run(slice_text).underline = True
                    paragraph_para.add_run(' ').underline = True
                else:
                    paragraph_para.add_run(slice_text)
        else:
            img_path = question['src']
            question_run = paragraph_para.add_run()
            question_run.add_picture(img_path)

    def write_attach(paragraph_para, content_para):
        attach = content_para.get('attach', None)
        if attach:
            if len(attach) == 1:
                for attach_img_path in attach:
                    attach_run = paragraph_para.add_run()
                    attach_run.add_picture(attach_img_path)
            else:
                attach_count = len(attach)
                for m, attach_img_path in enumerate(attach):
                    attach_run = paragraph_para.add_run()
                    attach_run.add_picture(attach_img_path)
                    if m+1 == attach_count:
                        ...
                    else:
                        paragraph_para.add_run('\n')

    def write_remark(paragraph_para, content_para):
        remark = content_para.get('remark', None)
        if remark:
            if remark == '':
                pass
            else:
                paragraph_para.add_run('\n')
                paragraph_para.add_run(remark)

    document = Document()
    document.styles['Normal'].font.name = doc_font_name_para
    document.styles['Normal'].font.size = doc_font_size_para
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), doc_font_name_para)
    text_dict_length = len(text_dict_para)
    for i in range(1, text_dict_length+1):
        div = text_dict_para[str(i)]
        content = div['content']
        if div['type'] == 'title':
            document.add_heading(content, 1)
        elif div['type'] == 'tag':
            document.add_heading(content, 3)
        elif div['type'] == '选择题':
            optional_question_p = document.add_paragraph()
            optional_question_p.paragraph_format.space_after = paragraph_apace_after_para
            write_question(optional_question_p, content)
            optional_question_p.add_run('\n')
            options = content['options']
            options_img_path = options.get('img', None)
            if options_img_path:
                options_img_run = optional_question_p.add_run()
                options_img_run.add_picture(options_img_path)
                optional_question_p.add_run('\n')
            else:
                items = options.get('items')
                item_max_length = 0
                for item in items:
                    current_item_length = len(item)
                    if current_item_length > item_max_length:
                        item_max_length = current_item_length
                items_length = len(items)
                if item_max_length < option_length_para:
                    for j, item in enumerate(items):
                        if j+1 == items_length:
                            if (j + 1) % 2 == 1:
                                padding_count = option_length_para - len(item)
                                optional_question_p.add_run(item + ' ' * padding_count)
                            else:
                                optional_question_p.add_run(item)
                        else:
                            if (j + 1) % 2 == 1:
                                padding_count = option_length_para - len(item)
                                optional_question_p.add_run(item + ' ' * padding_count)
                            else:
                                optional_question_p.add_run(item + '\n')
                else:
                    for k, item in enumerate(items):
                        if k+1 == items_length:
                            optional_question_p.add_run(item)
                        else:
                            optional_question_p.add_run(item+'\n')
            write_attach(optional_question_p, content)
            write_remark(optional_question_p, content)
        elif div['type'] == '填空题':
            fill_blank_question_p = document.add_paragraph()
            fill_blank_question_p.paragraph_format.space_after = paragraph_apace_after_para
            write_question(fill_blank_question_p, content)
            if content.get('attach', None):
                fill_blank_question_p.add_run('\n')
            write_attach(fill_blank_question_p, content)
            write_remark(fill_blank_question_p, content)
        elif div['type'] == '简答题':
            summary_question_p = document.add_paragraph()
            summary_question_p.paragraph_format.space_after = paragraph_apace_after_para
            write_question(summary_question_p, content)
            summary_question_p.add_run('\n')
            answer = content['answer']
            if answer['type'] == 'text':
                summary_question_p.add_run(answer['content']+'\n')
            else:
                answer_img_path = answer['src']
                answer_img_run = summary_question_p.add_run()
                answer_img_run.add_picture(answer_img_path)
                summary_question_p.add_run('\n')
            write_attach(summary_question_p, content)
            write_remark(summary_question_p, content)
    document.save(docx_filename)


if __name__ == '__main__':
    text_dict = read_from_json_file(json_filename)
    pprint.pprint(text_dict)
    write_in_docx(text_dict, docx_path)
